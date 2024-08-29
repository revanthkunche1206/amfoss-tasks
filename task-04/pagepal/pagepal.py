import os
import csv
import tempfile
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.constants import ParseMode
from telegram.ext import ApplicationBuilder, CommandHandler, CallbackQueryHandler, MessageHandler, filters
from googleapiclient.discovery import build
from docx import Document
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv

load_dotenv()

BOT_TOKEN = os.getenv("BOT_TOKEN")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

service = build('books', 'v1', developerKey=GOOGLE_API_KEY, cache_discovery=False)

reading_list = []

async def start(update: Update, context):
    await update.message.reply_text("Welcome to PagePal! I can help you find books by genre, provide previews, and manage your reading list.")

async def help_command(update: Update, context):
    help_text = (
        "/start - Welcome message\n"
        "/book - Find books by genre\n"
        "/preview - Get a preview link of a book\n"
        "/list - Manage your reading list\n"
        "/reading_list - View your reading list\n"
        "/help - Show this help message"
    )
    await update.message.reply_text(help_text)

async def book(update: Update, context):
    await update.message.reply_text("Please enter the genre you want to explore:")
    context.user_data['action'] = 'genre_search'

async def handle_genre_search(update: Update, context):
    genre = update.message.text
    results = service.volumes().list(q=f"subject:{genre}", maxResults=5).execute()
    if not results.get('items'):
        await update.message.reply_text("No books found for this genre.")
        return


    with tempfile.NamedTemporaryFile(delete=False, mode='w', newline='', suffix=".csv") as temp_csv:
        writer = csv.writer(temp_csv)
        writer.writerow(["Title", "Author", "Description", "Published Year", "Language", "Preview Link"])

        for item in results['items']:
            volume_info = item['volumeInfo']
            title = volume_info.get('title', 'N/A')
            authors = ', '.join(volume_info.get('authors', ['Unknown']))
            description = volume_info.get('description', 'N/A').replace("\n", " ")
            published_date = volume_info.get('publishedDate', 'N/A')
            language = volume_info.get('language', 'N/A')
            preview_link = volume_info.get('previewLink', 'N/A')
            writer.writerow([title, authors, description, published_date, language, preview_link])

        temp_csv.flush()
        await context.bot.send_document(chat_id=update.message.chat_id, document=open(temp_csv.name, 'rb'))


    context.user_data['action'] = None

async def preview(update: Update, context):
    await update.message.reply_text("Please enter the book name you want a preview for:")
    context.user_data['action'] = 'preview_search'

async def handle_preview_search(update: Update, context):
    book_name = update.message.text
    results = service.volumes().list(q=f"intitle:{book_name}", maxResults=1).execute()

    if not results.get('items'):
        await update.message.reply_text("No preview available for this book.")
        return

    preview_link = results['items'][0]['volumeInfo'].get('previewLink', 'No preview link available')

    await update.message.reply_text(
        f"[Click here to preview the book]({preview_link})",
        parse_mode=ParseMode.MARKDOWN
    )


    context.user_data['action'] = None

async def list_command(update: Update, context):
    await update.message.reply_text("Please type in the book name to manage and then execute /reading_list")

async def handle_button(update: Update, context):
    query = update.callback_query
    await query.answer()

    if query.data == 'add':
        await query.edit_message_text("Please enter the book title and preview link (comma-separated):")
        context.user_data['action'] = 'add'
    elif query.data == 'delete':
        await query.edit_message_text("Please enter the book title you want to delete:")
        context.user_data['action'] = 'delete'
    elif query.data == 'view':
        if not reading_list:
            await query.edit_message_text("Your reading list is empty.")
        else:
            doc = Document()
            doc.add_heading('Reading List', 0)

            for title, link in reading_list:
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(14)
                title_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

                if link != 'N/A':
                    link_paragraph = doc.add_paragraph()
                    link_run = link_paragraph.add_run(link)
                    link_run.italic = True
                    link_run.font.size = Pt(12)
                    link_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                doc.add_paragraph()

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                doc.save(temp_docx.name)
                await context.bot.send_document(chat_id=query.message.chat_id, document=open(temp_docx.name, 'rb'))

async def handle_reading_list_action(update: Update, context):
    if context.user_data.get('action') == 'add':
        book_info = update.message.text.split(',')
        if len(book_info) != 2:
            await update.message.reply_text("Invalid format. Please enter as 'Title, Preview Link'.")
            return
        title, link = book_info
        reading_list.append((title.strip(), link.strip()))
        await update.message.reply_text(f"Added {title} to your reading list.")
    elif context.user_data.get('action') == 'delete':
        title = update.message.text.strip()
        for i, (book_title, _) in enumerate(reading_list):
            if book_title.lower() == title.lower():
                del reading_list[i]
                await update.message.reply_text(f"Deleted {title} from your reading list.")
                break
        else:
            await update.message.reply_text(f"{title} not found in your reading list.")
    context.user_data['action'] = None

async def reading_list(update: Update, context):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add')],
        [InlineKeyboardButton("Delete a book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Choose an option:", reply_markup=reply_markup)

if __name__ == '__main__':
    application = ApplicationBuilder().token(BOT_TOKEN).build()

    application.add_handler(CommandHandler('start', start))
    application.add_handler(CommandHandler('help', help_command))
    application.add_handler(CommandHandler('book', book))
    application.add_handler(CommandHandler('preview', preview))
    application.add_handler(CommandHandler('list', list_command))
    application.add_handler(CommandHandler('reading_list', reading_list))
    application.add_handler(CallbackQueryHandler(handle_button))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_genre_search))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_preview_search))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_reading_list_action))

    application.run_polling()
